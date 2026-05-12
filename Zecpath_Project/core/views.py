from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import generics
from rest_framework.permissions import AllowAny
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.views import TokenObtainPairView
from .models import Job, CustomUser, Application ,Employer , AdminLog
from .serializers import JobSerializer, UserSerializer , CandidateSerializer, EmployerSerializer
from .permissions import IsAdmin, IsEmployer, IsCandidate
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.filters import SearchFilter
from django_filters.rest_framework import DjangoFilterBackend
from rest_framework.pagination import PageNumberPagination
from .serializers import ApplicationSerializer
from .permissions import IsCandidate
from rest_framework.generics import ListAPIView
from rest_framework.filters import SearchFilter
import re
import PyPDF2
import pdfplumber
from docx import Document
from .serializers import ResumeUploadSerializer
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser




SKILLS_LIBRARY = [
    "python",
    "django",
    "rest api",
    "sql",
    "mysql",
    "html",
    "css",
    "javascript",
    "react",
    "java",
    "git",
    "docker"
]



class JobPagination(PageNumberPagination):
    page_size = 5



class SignupAPI(generics.CreateAPIView):
    queryset = CustomUser.objects.all()
    serializer_class = UserSerializer
    permission_classes = [AllowAny]

    def create(self, request, *args, **kwargs):
        email = request.data.get("email")

        # ✅ check email exists in request
        if not email:
            return Response({"error": "Email is required"}, status=400)

        # ✅ check duplicate properly
        if CustomUser.objects.filter(email=email).exists():
            return Response({"error": "Email already exists"}, status=400)

        return super().create(request, *args, **kwargs)



# ✅ Login
class LoginAPI(TokenObtainPairView):
    pass






class JobCreateAPI(APIView):
    permission_classes = [IsAuthenticated, IsEmployer]

    def post(self, request):
        try:
            employer = request.user.employer
        except:
            return Response({"error": "Employer profile not found"}, status=400)

        serializer = JobSerializer(data=request.data)

        if serializer.is_valid():
            serializer.save(employer=employer)
            return Response(serializer.data, status=201)

        return Response(serializer.errors, status=400)








class JobUpdateAPI(APIView):
    permission_classes = [IsAuthenticated, IsEmployer]

    def put(self, request, pk):
        employer = request.user.employer

        try:
            job = Job.objects.get(id=pk, employer=employer)
        except Job.DoesNotExist:
            return Response({"error": "Not allowed"}, status=403)

        serializer = JobSerializer(job, data=request.data, partial=True)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)

        return Response(serializer.errors, status=400)









class JobToggleAPI(APIView):
    permission_classes = [IsAuthenticated, IsEmployer]

    def post(self, request, pk):
        employer = request.user.employer

        try:
            job = Job.objects.get(id=pk, employer=employer)
        except Job.DoesNotExist:
            return Response({"error": "Not allowed"}, status=403)

        job.status = 'inactive' if job.status == 'active' else 'active'
        job.save()

        return Response({"status": job.status})








class JobListAPI(generics.ListAPIView):
    queryset = Job.objects.filter(status='active').select_related('employer','employer__user')
    serializer_class = JobSerializer
    permission_classes = [AllowAny]

    pagination_class = JobPagination

    filter_backends = [DjangoFilterBackend, SearchFilter]

    filterset_fields = {
        'skills': ['exact'],
        'location': ['exact'],
        'job_type': ['exact'],
        'experience': ['gte', 'lte'],
        'salary_min': ['gte'],
        'salary_max': ['lte'],
    }

    search_fields = ['title', 'description', 'skills']






class ApplyJobAPI(APIView):
    permission_classes = [IsAuthenticated, IsCandidate]

    def post(self, request):
        job_id = request.data.get('job')

        try:
            candidate = request.user.candidate
        except:
            return Response({"error": "Candidate profile not found"}, status=400)

        try:
            job = Job.objects.get(id=job_id, status='active')
        except Job.DoesNotExist:
            return Response({"error": "Job not available"}, status=404)

        # ❌ Duplicate prevention
        if Application.objects.filter(job=job, candidate=candidate).exists():
            return Response({"error": "Already applied"}, status=400)

        # ✅ Resume binding
        resume = candidate.resume

        application = Application.objects.create(
            job=job,
            candidate=candidate,
            resume=resume
        )

        return Response({
            "message": "Applied successfully",
            "application_id": application.id
        }, status=201)






class UserTestAPI(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def get(self, request):
        users = CustomUser.objects.all()
        serializer = UserSerializer(users, many=True)
        return Response(serializer.data)
    





class CandidateProfileAPI(APIView):
    permission_classes = [IsAuthenticated, IsCandidate]
    parser_classes = [MultiPartParser, FormParser,JSONParser]   # 👈 ADD THIS LINE

    def get(self, request):
        candidate = request.user.candidate
        return Response(CandidateSerializer(candidate).data)



    def put(self, request):

        candidate = request.user.candidate

        data = request.data.copy()

    # ✅ safe conversion
        if data.get('experience'):
            data['experience'] = float(data.get('experience'))

        serializer = CandidateSerializer(
            candidate,
            data=data,
            partial=True
        )

        if serializer.is_valid():

            serializer.save()

            return Response(serializer.data)

        return Response(serializer.errors, status=400)


    def delete(self, request):
        candidate = request.user.candidate
        candidate.is_active = False
        candidate.save()
        return Response({"message": "Profile soft deleted"})
    






class EmployerProfileAPI(APIView):
    permission_classes = [IsAuthenticated, IsEmployer]

    def get(self, request):
        employer = request.user.employer
        return Response(EmployerSerializer(employer).data)

    def put(self, request):
        employer = request.user.employer
        serializer = EmployerSerializer(employer, data=request.data, partial=True)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)

        return Response(serializer.errors, status=400)

    def delete(self, request):
        employer = request.user.employer
        employer.is_active = False
        employer.save()
        return Response({"message": "Profile soft deleted"})
    

    


class MyApplicationsAPI(generics.ListAPIView):
    serializer_class = ApplicationSerializer
    permission_classes = [IsAuthenticated, IsCandidate]

    def get_queryset(self):
        return Application.objects.filter(
            candidate=self.request.user.candidate
        ).select_related('job')
    




class UpdateApplicationStatusAPI(APIView):
    permission_classes = [IsAuthenticated, IsEmployer]

    def post(self, request, pk):
        employer = request.user.employer

        try:
            application = Application.objects.select_related('job').get(id=pk)
        except Application.DoesNotExist:
            return Response({"error": "Application not found"}, status=404)

        # ✅ Ownership check
        if application.job.employer != employer:
            return Response({"error": "Not allowed"}, status=403)

        new_status = request.data.get("status")

        valid_transitions = {
            'applied': ['shortlisted', 'rejected'],
            'shortlisted': ['interview', 'rejected'],
            'interview': ['selected', 'rejected'],
            'selected': [],
            'rejected': []
        }

        current_status = application.status

        if new_status not in valid_transitions.get(current_status, []):
            return Response({
                "error": f"Cannot move from {current_status} → {new_status}"
            }, status=400)

        application.status = new_status
        application.save()

        return Response({
            "message": "Status updated",
            "new_status": application.status
        })
    




class EmployerJobsAPI(ListAPIView):
    serializer_class = JobSerializer
    permission_classes = [IsAuthenticated, IsEmployer]

    def get_queryset(self):
        return Job.objects.filter(
            employer=self.request.user.employer
        )
    










class JobApplicantsAPI(ListAPIView):
    serializer_class = ApplicationSerializer
    permission_classes = [IsAuthenticated, IsEmployer]

    filter_backends = [SearchFilter]
    search_fields = ['candidate__user__email']

    def get_queryset(self):
        job_id = self.kwargs['job_id']
        return Application.objects.filter(
            job__id=job_id,
            job__employer=self.request.user.employer
        )
    






class EmployerAnalyticsAPI(APIView):
    permission_classes = [IsAuthenticated, IsEmployer]

    def get(self, request):
        employer = request.user.employer

        total_jobs = Job.objects.filter(employer=employer).count()
        total_applications = Application.objects.filter(job__employer=employer).count()
        shortlisted = Application.objects.filter(
            job__employer=employer,
            status='shortlisted'
        ).count()

        return Response({
            "total_jobs": total_jobs,
            "total_applications": total_applications,
            "shortlisted": shortlisted
        })
    









class AppliedJobsAPI(generics.ListAPIView):
    serializer_class = ApplicationSerializer
    permission_classes = [IsAuthenticated, IsCandidate]

    def get_queryset(self):
        return Application.objects.filter(
            candidate=self.request.user.candidate
        ).select_related('job')
    

    




class RecommendedJobsAPI(generics.ListAPIView):
    serializer_class = JobSerializer
    permission_classes = [IsAuthenticated, IsCandidate]

    def get_queryset(self):
        candidate = self.request.user.candidate
        skills = candidate.skills

        return Job.objects.filter(
            skills__icontains=skills,
            status='active'
        )
    





class ApproveEmployerAPI(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def post(self, request, pk):
        try:
            employer = Employer.objects.get(id=pk)
        except Employer.DoesNotExist:
            return Response({"error": "Employer not found"}, status=404)

        employer.is_verified = True
        employer.save()

        return Response({
            "message": "Employer approved"
        })
    






class BlockUserAPI(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def post(self, request, pk):
        try:
            user = CustomUser.objects.get(id=pk)
        except CustomUser.DoesNotExist:
            return Response({"error": "User not found"}, status=404)

        user.is_active = False
        user.save()

        # ✅ Admin audit log
        AdminLog.objects.create(
            admin=request.user,
            action=f"Blocked user {user.email}"
        )

        return Response({
            "message": "User blocked"
        })





class RemoveJobAPI(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def delete(self, request, pk):
        try:
            job = Job.objects.get(id=pk)
        except Job.DoesNotExist:
            return Response({"error": "Job not found"}, status=404)

        job.delete()

        return Response({
            "message": "Spam job removed"
        })
    





class PlatformStatsAPI(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def get(self, request):

        total_users = CustomUser.objects.count()
        total_jobs = Job.objects.count()
        total_applications = Application.objects.count()

        employers = CustomUser.objects.filter(role='employer').count()
        candidates = CustomUser.objects.filter(role='candidate').count()

        return Response({
            "total_users": total_users,
            "total_jobs": total_jobs,
            "total_applications": total_applications,
            "employers": employers,
            "candidates": candidates
        })
    






class AdminLogsAPI(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def get(self, request):

        logs = AdminLog.objects.all().values(
            'id',
            'action',
            'created_at'
        )

        return Response(logs)
    







class ResumeParserAPI(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        serializer = ResumeUploadSerializer(data=request.data)

        if serializer.is_valid():
            file = serializer.validated_data['file']

            extracted_text = ""

            # ✅ PDF Extraction
            if file.name.endswith('.pdf'):

                try:
                    pdf = PyPDF2.PdfReader(file)

                    for page in pdf.pages:
                        extracted_text += page.extract_text()

                except:
                    return Response({
                        "error": "PDF extraction failed"
                    }, status=400)

            # ✅ DOCX Extraction
            elif file.name.endswith('.docx'):

                try:
                    doc = Document(file)

                    for para in doc.paragraphs:
                        extracted_text += para.text + "\n"

                except:
                    return Response({
                        "error": "DOCX extraction failed"
                    }, status=400)

            else:
                return Response({
                    "error": "Only PDF or DOCX supported"
                }, status=400)

            # ✅ CLEANING
            cleaned_text = re.sub(r'\s+', ' ', extracted_text).strip()

            return Response({
                "raw_text": extracted_text,
                "cleaned_text": cleaned_text
            })

        return Response(serializer.errors, status=400)
    






class ResumeNLPAPI(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):

        serializer = ResumeUploadSerializer(data=request.data)

        if serializer.is_valid():

            file = serializer.validated_data['file']

            extracted_text = ""

            # ✅ PDF extraction
            if file.name.endswith('.pdf'):

                try:
                    pdf = PyPDF2.PdfReader(file)

                    for page in pdf.pages:
                        text = page.extract_text()

                        if text:
                            extracted_text += text

                except:
                    return Response({
                        "error": "PDF extraction failed"
                    }, status=400)

            # ✅ DOCX extraction
            elif file.name.endswith('.docx'):

                try:
                    doc = Document(file)

                    for para in doc.paragraphs:
                        extracted_text += para.text + "\n"

                except:
                    return Response({
                        "error": "DOCX extraction failed"
                    }, status=400)

            else:
                return Response({
                    "error": "Only PDF or DOCX supported"
                }, status=400)

            # ✅ CLEAN TEXT
            cleaned_text = re.sub(r'\s+', ' ', extracted_text).strip()

            # ✅ TOKENIZATION
            tokens = cleaned_text.lower().split()

            # ✅ SKILL EXTRACTION
            found_skills = []

            for skill in SKILLS_LIBRARY:

                if skill.lower() in cleaned_text.lower():
                    found_skills.append(skill)

            # ✅ EXPERIENCE EXTRACTION
            experience = re.findall(
                r'(\d+)\s+years?',
                cleaned_text.lower()
            )

            # ✅ EDUCATION DETECTION
            education_keywords = [
                "btech",
                "bca",
                "mca",
                "bsc",
                "msc",
                "computer science"
            ]

            education_found = []

            for edu in education_keywords:

                if edu.lower() in cleaned_text.lower():
                    education_found.append(edu)

            # ✅ STRUCTURED JSON OUTPUT
            parsed_resume = {
                "skills": found_skills,
                "experience_years": experience,
                "education": education_found,
                "tokens_count": len(tokens)
            }

            return Response({
                "cleaned_text": cleaned_text,
                "parsed_resume": parsed_resume
            })

        return Response(serializer.errors, status=400)
    





class ATSMatchAPI(APIView):
    permission_classes = [IsAuthenticated, IsEmployer]

    def get(self, request, job_id):

        try:
            job = Job.objects.get(
                id=job_id,
                employer=request.user.employer
            )

        except Job.DoesNotExist:
            return Response({
                "error": "Job not found"
            }, status=404)

        applications = Application.objects.filter(job=job)

        ranked_candidates = []

        for app in applications:

            candidate = app.candidate

            score = 0

            # ✅ CLEAN SKILLS
            job_skills = [
                skill.strip().lower()
                for skill in job.skills.split(',')
            ]

            candidate_skills = [
                skill.strip().lower()
                for skill in candidate.skills.split(',')
            ]

            print("JOB SKILLS:", job_skills)
            print("CANDIDATE SKILLS:", candidate_skills)

            matched_skills = []

            # ✅ SKILL MATCHING
            for skill in job_skills:

                if skill in candidate_skills:

                    matched_skills.append(skill)

                    score += 20

            print("MATCHED:", matched_skills)

            # ✅ EXPERIENCE MATCHING
            if candidate.experience:

                if candidate.experience >= job.experience:
                    score += 20

            # ✅ EDUCATION BONUS
            if candidate.education:
                score += 20

            # ✅ MAX SCORE LIMIT
            if score > 100:
                score = 100

            # ✅ SAVE SCORE
            app.ats_score = score
            app.save()

            ranked_candidates.append({
                "candidate_email": candidate.user.email,
                "matched_skills": matched_skills,
                "ats_score": score
            })

        ranked_candidates = sorted(
            ranked_candidates,
            key=lambda x: x['ats_score'],
            reverse=True
        )

        return Response({
            "job": job.title,
            "ranked_candidates": ranked_candidates
        })
    







class AutoShortlistAPI(APIView):
    permission_classes = [IsAuthenticated, IsEmployer]

    def post(self, request, job_id):

        try:
            job = Job.objects.get(
                id=job_id,
                employer=request.user.employer
            )

        except Job.DoesNotExist:
            return Response({
                "error": "Job not found"
            }, status=404)

        applications = Application.objects.filter(job=job)

        shortlisted = []
        rejected = []

        for app in applications:

            # ✅ Threshold logic
            if app.ats_score >= 60:

                app.status = "shortlisted"
                shortlisted.append(app.candidate.user.email)

            else:

                app.status = "rejected"
                rejected.append(app.candidate.user.email)

            app.save()

        return Response({
            "job": job.title,
            "shortlisted_candidates": shortlisted,
            "rejected_candidates": rejected
        })